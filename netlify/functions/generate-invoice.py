"""
Netlify Function: Generate Invoice
This function generates a DOCX invoice from the provided data using a template.
"""

import json
import base64
import io
import os
import re
import tempfile
import zipfile
from datetime import datetime
from urllib.request import urlopen, Request
from urllib.error import URLError

# Try to import python-docx
try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls, qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# ================== CONFIGURATION ==================

# Template URL - hosted on a public URL
# You can host this on GitHub raw, Google Drive, or any public URL
TEMPLATE_URL = os.environ.get('INVOICE_TEMPLATE_URL', '')

# Paid stamp and signature URLs (optional)
PAID_STAMP_URL = os.environ.get('PAID_STAMP_URL', 'https://drive.google.com/uc?export=download&id=1W9PL0DtP0TUk7IcGiMD_ZuLddtQ8gjNo')
SIGNATURE_URL = os.environ.get('SIGNATURE_URL', 'https://drive.google.com/uc?export=download&id=1b6Dcg4spQmvLUMd4neBtLNfdr5l7QtPJ')

# ================== HELPER FUNCTIONS ==================

def format_currency(amount):
    """Format amount as Indonesian Rupiah"""
    if amount == 0:
        return ""
    elif amount == int(amount):
        return f"Rp {int(amount):,}".replace(',', '.')
    else:
        return f"Rp {amount:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')

def sanitize_filename(name):
    """Remove or replace characters that are invalid in file names"""
    return re.sub(r'[<>:"/\\|?*]', '_', name).replace(' ', '_')

def set_cell_border(cell, side, color="FFFFFF", sz=4):
    """Set border for a cell"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    side_mapping = {'top': 'top', 'bottom': 'bottom', 'left': 'left', 'right': 'right'}
    border_name = side_mapping.get(side.lower())
    if border_name:
        border = parse_xml(f'<w:{border_name} {nsdecls("w")} w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>')
        tcBorders = tcPr.first_child_found_in("w:tcBorders")
        if tcBorders is None:
            tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
            tcPr.append(tcBorders)
        tcBorders.append(border)

def set_white_borders(cell, sz=4):
    """Set white borders on all sides of a cell"""
    for border in ['top', 'bottom', 'left', 'right']:
        set_cell_border(cell, border, color="FFFFFF", sz=sz)

def set_cell_font(cell, font_name="Courier New", font_size=10):
    """Set font for all text in a cell"""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = font_name
            run.font.size = Pt(font_size)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def apply_cell_style(cell, bg_color="ddefd5"):
    """Apply styling to a cell"""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg_color}" />')
    cell._tc.get_or_add_tcPr().append(shading_elm)
    set_white_borders(cell, sz=6)
    set_cell_font(cell)

def replace_placeholders(doc, replacements):
    """Replace placeholders in the document"""
    for paragraph in doc.paragraphs:
        for key, value in replacements.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, str(value))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in replacements.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, str(value))
    return doc

def update_items_table(doc, items):
    """Update the items table in the document"""
    items_table = doc.tables[0]
    
    # Set borders for existing rows
    for i in range(len(items_table.rows)):
        for cell in items_table.rows[i].cells:
            set_white_borders(cell, sz=6)
    
    # Remove all rows except header and placeholder
    while len(items_table.rows) > 2:
        items_table._tbl.remove(items_table.rows[2]._tr)
    
    placeholder_row = items_table.rows[1]
    
    # Add items
    for item in items:
        row = items_table.add_row()
        row.cells[0].text = item['description']
        row.cells[1].text = format_currency(item['unit_price'])
        quantity = item['quantity']
        if quantity == int(quantity):
            row.cells[2].text = str(int(quantity))
        else:
            row.cells[2].text = str(quantity)
        row.cells[3].text = format_currency(item['total'])
        
        # Apply styling
        alignments = [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.RIGHT, 
                     WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.RIGHT]
        for i, cell in enumerate(row.cells):
            apply_cell_style(cell)
            for paragraph in cell.paragraphs:
                paragraph.alignment = alignments[i]
    
    # Remove placeholder row
    items_table._tbl.remove(placeholder_row._tr)
    
    return doc

def style_financial_table(doc, apply_late_fee):
    """Style the financial summary table"""
    financial_table = doc.tables[1]
    for row in financial_table.rows:
        for cell in row.cells:
            set_white_borders(cell)
            set_cell_font(cell)
        for paragraph in row.cells[1].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    if apply_late_fee:
        late_fee_cell = financial_table.rows[3].cells[0]
        if "LATE FEE" in late_fee_cell.text:
            original_text = late_fee_cell.text
            late_fee_cell.text = ""
            paragraph = late_fee_cell.paragraphs[0]
            run = paragraph.add_run(original_text)
            run.font.color.rgb = RGBColor.from_string('d95132')
            run.font.name = "Courier New"
            run._element.rPr.rFonts.set(qn('w:eastAsia'), "Courier New")

def download_template():
    """Download the invoice template from URL"""
    if not TEMPLATE_URL:
        raise ValueError("INVOICE_TEMPLATE_URL environment variable not set")
    
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
    }
    request = Request(TEMPLATE_URL, headers=headers)
    
    try:
        with urlopen(request, timeout=30) as response:
            return io.BytesIO(response.read())
    except URLError as e:
        raise Exception(f"Failed to download template: {str(e)}")

def create_embedded_template():
    """
    Create a basic invoice template if no template URL is provided.
    This is a fallback option.
    """
    doc = Document()
    
    # Add company header
    header = doc.add_paragraph()
    header_run = header.add_run("INVOICE")
    header_run.bold = True
    header_run.font.size = Pt(24)
    
    # Add placeholder paragraphs
    doc.add_paragraph("From: Marketix Lab")
    doc.add_paragraph("+6281316198226 | marketixlab@gmail.com")
    doc.add_paragraph("")
    doc.add_paragraph("Issued to:")
    doc.add_paragraph("{{client_name}}")
    doc.add_paragraph("{{client_phone}}")
    doc.add_paragraph("{{client_email}}")
    doc.add_paragraph("{{client_address}}")
    doc.add_paragraph("")
    doc.add_paragraph("Invoice No: {{invoice_number}}")
    doc.add_paragraph("Date: {{invoice_date}}")
    doc.add_paragraph("Due Date: {{due_date}}")
    doc.add_paragraph("")
    
    # Create items table
    items_table = doc.add_table(rows=2, cols=4)
    items_table.style = 'Table Grid'
    
    # Header row
    header_cells = items_table.rows[0].cells
    header_cells[0].text = "DESCRIPTION"
    header_cells[1].text = "UNIT PRICE"
    header_cells[2].text = "QUANTITY"
    header_cells[3].text = "TOTAL"
    
    # Placeholder row
    placeholder_cells = items_table.rows[1].cells
    placeholder_cells[0].text = "{{service_description}}"
    placeholder_cells[1].text = "{{unit_price}}"
    placeholder_cells[2].text = "{{quantity}}"
    placeholder_cells[3].text = "{{total}}"
    
    doc.add_paragraph("")
    
    # Create financial table
    financial_table = doc.add_table(rows=5, cols=2)
    financial_table.style = 'Table Grid'
    
    rows_data = [
        ("SUBTOTAL", "[subtotal]"),
        ("TAX", "[tax]"),
        ("DISCOUNT", "[discount]"),
        ("{{LATE FEE:}}", "[latefee]"),
        ("GRAND TOTAL", "[grandtotal]")
    ]
    
    for i, (label, value) in enumerate(rows_data):
        financial_table.rows[i].cells[0].text = label
        financial_table.rows[i].cells[1].text = value
    
    doc.add_paragraph("")
    doc.add_paragraph("Bank Details:")
    doc.add_paragraph("Bank Name: Bank Mandiri")
    doc.add_paragraph("Account Name: ARIELLA SHEEHAN EVAN")
    doc.add_paragraph("Account No: 1760002641759")
    
    return doc

def generate_invoice(data):
    """Generate the invoice document"""
    # Try to download template, fall back to embedded template
    try:
        if TEMPLATE_URL:
            template_data = download_template()
            doc = Document(template_data)
        else:
            doc = create_embedded_template()
    except Exception as e:
        print(f"Warning: Could not load template ({str(e)}), using embedded template")
        doc = create_embedded_template()
    
    # Prepare replacements
    replacements = {
        **data.get('client_info', {}),
        **data.get('invoice_details', {}),
        **data.get('financials', {})
    }
    
    # Handle late fee
    if data.get('apply_late_fee', False):
        replacements['{{LATE FEE:}}'] = 'LATE FEE'
    else:
        replacements['{{LATE FEE:}}'] = ''
        replacements['[latefee]'] = ''
    
    # Replace placeholders
    doc = replace_placeholders(doc, replacements)
    
    # Update items table
    items = data.get('items', [])
    if items:
        doc = update_items_table(doc, items)
    
    # Style financial table
    try:
        style_financial_table(doc, data.get('apply_late_fee', False))
    except Exception as e:
        print(f"Warning: Could not style financial table: {str(e)}")
    
    # Apply font to all paragraphs
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = "Courier New"
            run._element.rPr.rFonts.set(qn('w:eastAsia'), "Courier New")
    
    # Generate output
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    
    # Generate filename
    client_name = data.get('client_info', {}).get('{{client_name}}', 'Client')
    client_name = sanitize_filename(client_name)
    invoice_number = data.get('invoice_number', 'INV')
    prefix = "Paid_Invoice" if data.get('mark_as_paid', False) else "Invoice"
    filename = f"{prefix}_{invoice_number}_{client_name}.docx"
    
    return output.getvalue(), filename

# ================== NETLIFY HANDLER ==================

def handler(event, context):
    """Netlify Function handler"""
    
    # Handle CORS preflight
    if event.get('httpMethod') == 'OPTIONS':
        return {
            'statusCode': 200,
            'headers': {
                'Access-Control-Allow-Origin': '*',
                'Access-Control-Allow-Headers': 'Content-Type',
                'Access-Control-Allow-Methods': 'POST, OPTIONS'
            },
            'body': ''
        }
    
    # Only accept POST requests
    if event.get('httpMethod') != 'POST':
        return {
            'statusCode': 405,
            'headers': {
                'Access-Control-Allow-Origin': '*',
                'Content-Type': 'application/json'
            },
            'body': json.dumps({'error': 'Method not allowed'})
        }
    
    # Check if python-docx is available
    if not DOCX_AVAILABLE:
        return {
            'statusCode': 500,
            'headers': {
                'Access-Control-Allow-Origin': '*',
                'Content-Type': 'application/json'
            },
            'body': json.dumps({'error': 'python-docx library not available'})
        }
    
    try:
        # Parse request body
        body = event.get('body', '{}')
        if event.get('isBase64Encoded', False):
            body = base64.b64decode(body).decode('utf-8')
        
        data = json.loads(body)
        
        # Validate required fields
        required_fields = ['client_info', 'invoice_details', 'items', 'financials']
        for field in required_fields:
            if field not in data:
                return {
                    'statusCode': 400,
                    'headers': {
                        'Access-Control-Allow-Origin': '*',
                        'Content-Type': 'application/json'
                    },
                    'body': json.dumps({'error': f'Missing required field: {field}'})
                }
        
        # Generate invoice
        docx_bytes, filename = generate_invoice(data)
        
        # Return the DOCX file as base64
        return {
            'statusCode': 200,
            'headers': {
                'Access-Control-Allow-Origin': '*',
                'Content-Type': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                'Content-Disposition': f'attachment; filename="{filename}"'
            },
            'body': base64.b64encode(docx_bytes).decode('utf-8'),
            'isBase64Encoded': True
        }
        
    except json.JSONDecodeError:
        return {
            'statusCode': 400,
            'headers': {
                'Access-Control-Allow-Origin': '*',
                'Content-Type': 'application/json'
            },
            'body': json.dumps({'error': 'Invalid JSON in request body'})
        }
    except Exception as e:
        return {
            'statusCode': 500,
            'headers': {
                'Access-Control-Allow-Origin': '*',
                'Content-Type': 'application/json'
            },
            'body': json.dumps({'error': str(e)})
        }
